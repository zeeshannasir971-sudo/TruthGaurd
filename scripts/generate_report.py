from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def main(output_path="FakeNews_Project_Report.docx"):
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("Fake News Detection with Sentiment Analysis — Project Report")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "Date: 2026-01-08")

    add_heading(doc, "Introduction", level=1)
    add_paragraph(doc, "This project presents a practical system for detecting fake news by combining text-based machine learning features (TF‑IDF at word and character levels) with sentiment analysis. The system exposes a simple web interface to analyze article URLs or pasted text and uses web search to find corroborating sources.")

    add_heading(doc, "Background", level=1)
    add_paragraph(doc, "Online misinformation remains pervasive, impacting public discourse, trust, and decision-making. Traditional detectors focus on lexical features; however, emotionally manipulative content often carries distinct sentiment and stylistic signals. This project integrates sentiment and simple style features to complement TF‑IDF, aiming for stronger generalization.")

    add_heading(doc, "Motivation", level=1)
    add_paragraph(doc, "We developed this system to provide a fast, explainable baseline capable of high accuracy on common datasets while remaining transparent. The goal is to balance performance with interpretability, using standard models and features that are easy to audit and deploy.")

    add_heading(doc, "Gap Analysis", level=1)
    add_bullet(doc, "Many detectors rely solely on lexical TF‑IDF, missing emotional cues.")
    add_bullet(doc, "End‑to‑end tools often lack a clean web interface and URL extraction.")
    add_bullet(doc, "Few baselines include corroboration via external web sources.")

    add_heading(doc, "Objectives", level=1)
    add_bullet(doc, "Build a robust baseline using TF‑IDF (word + character).")
    add_bullet(doc, "Augment with sentiment and simple style features.")
    add_bullet(doc, "Provide a Flask web UI for URL and text analysis.")
    add_bullet(doc, "Offer lightweight web corroboration through search.")

    add_heading(doc, "System Overview", level=1)
    add_paragraph(doc, "The system consists of data preprocessing, feature engineering, model training, prediction, content extraction, corroboration, and a web interface.")
    add_bullet(doc, "Entry point: `app.py` registers web routes and serves the UI.")
    add_bullet(doc, "Routes: `src/web/routes.py` provides `/` (UI), `/analyze`, `/train`.")
    add_bullet(doc, "Model pipeline: `src/ml/pipeline.py` trains and predicts.")
    add_bullet(doc, "Utilities: preprocessing, sentiment, fetch, search modules under `src/utils/`.")

    add_heading(doc, "Modules", level=1)
    add_paragraph(doc, "Preprocessing (`src/utils/preprocess.py`): cleans text by removing URLs and punctuation, lowercasing, and collapsing whitespace. It also tokenizes and filters stopwords.")
    add_paragraph(doc, "Sentiment (`src/utils/sentiment.py`): uses TextBlob to compute polarity and subjectivity.")
    add_paragraph(doc, "Feature Engineering (`src/ml/pipeline.py`): builds TF‑IDF features at word and character levels, then adds sentiment and simple style features (exclamation ratio, uppercase ratio, punctuation density).")
    add_paragraph(doc, "Model Training (`src/ml/pipeline.py`): trains baseline and sentiment+style models, performs 80/20 train/test splits, and saves artifacts.")
    add_paragraph(doc, "Prediction (`src/ml/pipeline.py`): loads artifacts, generates features for input text, and outputs label, probability, and sentiment.")
    add_paragraph(doc, "Content Extraction (`src/utils/fetch.py`): extracts main article text via Trafilatura with BeautifulSoup fallback.")
    add_paragraph(doc, "Web Corroboration (`src/utils/search.py`): queries DuckDuckGo for related sources.")
    add_paragraph(doc, "Web UI (`src/web/templates/index.html`): provides tabs for URL and pasted text analysis and displays results.")

    add_heading(doc, "Models Used", level=1)
    add_bullet(doc, "TF‑IDF Vectorizers: word‑level (1–2 n‑grams, max_features 100k) and character‑level (3–4 n‑grams, max_features 50k).")
    add_bullet(doc, "Baseline Classifiers: LinearSVC with probability calibration and CV‑tuned Logistic Regression (liblinear).")
    add_bullet(doc, "Final Model: Logistic Regression trained on full data with balanced class weights.")

    add_heading(doc, "Dataset", level=1)
    add_paragraph(doc, "Default training dataset: `src/data/kaggle_fake_real_combined.csv` with columns `text` and `label` (0/1). A small sample is also provided for demonstration in `src/data/sample_news.csv`.")

    add_heading(doc, "Training and Evaluation", level=1)
    add_paragraph(doc, "The pipeline performs a stratified train/test split with `test_size=0.2`, i.e., 80% training and 20% testing. Accuracy and classification reports are printed for both baseline and sentiment+style models.")
    add_bullet(doc, "Split confirmation: “Data Split: 80% Training, 20% Testing”.")
    add_bullet(doc, "Baseline accuracy: ≈ 99.9% (example shown in recent run).")
    add_bullet(doc, "Sentiment+style accuracy: ≈ 99.7% (example shown in recent run).")
    add_paragraph(doc, "Artifacts are saved under `src/ml/artifacts/` including `fake_news_model.pkl`, `tfidf_word_vectorizer.pkl`, and `tfidf_char_vectorizer.pkl`.")

    add_heading(doc, "Architecture Details and Code References", level=1)
    add_bullet(doc, "Flask app registration: `app.py:4-8`")
    add_bullet(doc, "Routes and endpoints: `src/web/routes.py:10-42`")
    add_bullet(doc, "Pipeline training entry: `src/ml/pipeline.py:33-169`")
    add_bullet(doc, "Prediction function: `src/ml/pipeline.py:179-200`")
    add_bullet(doc, "Preprocessing utilities: `src/utils/preprocess.py:1-36`")
    add_bullet(doc, "Sentiment extraction: `src/utils/sentiment.py:1-15`")
    add_bullet(doc, "Web content extraction: `src/utils/fetch.py:1-48`")
    add_bullet(doc, "Corroboration search: `src/utils/search.py:1-24`")

    add_heading(doc, "Deployment and Usage", level=1)
    add_paragraph(doc, "Run locally: create a virtual env, install requirements, and start `app.py`. Analyze URLs or pasted text through the UI; alternatively use `/train` to retrain the model.")

    add_heading(doc, "Limitations", level=1)
    add_bullet(doc, "High reported accuracy may reflect dataset characteristics; real‑world generalization can vary.")
    add_bullet(doc, "Sentiment features are simple and may not capture nuanced rhetoric.")
    add_bullet(doc, "Web corroboration relies on search API and result quality.")

    add_heading(doc, "Future Work", level=1)
    add_bullet(doc, "Incorporate domain and source credibility signals.")
    add_bullet(doc, "Add transformer embeddings with efficient inference.")
    add_bullet(doc, "Improve explainability via feature attribution.")
    add_bullet(doc, "Expand multilingual support and robustness checks.")

    add_heading(doc, "Conclusion", level=1)
    add_paragraph(doc, "This project delivers an end‑to‑end, explainable baseline for fake news detection that blends TF‑IDF features with sentiment and style signals, packaged with a practical web interface and corroboration. It provides a strong foundation for iterative improvements while remaining transparent and easy to maintain.")

    doc.save(output_path)

if __name__ == "__main__":
    main()
